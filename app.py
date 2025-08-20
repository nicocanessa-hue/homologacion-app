# -*- coding: utf-8 -*-
import re, io, os, json
import pandas as pd
import streamlit as st
import pdfplumber
from docx import Document
from io import BytesIO
from unidecode import unidecode
from rapidfuzz import fuzz

# ================== UI BASE ==================
st.set_page_config(page_title="Homologación PRO", page_icon="🧪", layout="wide")
st.title("🧪 Homologación de Materias Primas — PRO")

st.markdown("""
Sube una **Especificación** (PDF o DOCX) y **1..N documentos del Proveedor** (PDF/DOCX).
El sistema:
1) extrae **texto y tablas** (pdfplumber → Camelot → Tabula),
2) detecta **esquemas** Parametro/Min/Target/Max/Unidad,
3) normaliza **sinónimos ES/EN** y **unidades**,
4) **consolida** múltiples documentos por proveedor y
5) **compara** (Cumple / No cumple / No informado / Supera).
""")

# ================== UTILIDADES ==================
def nrm(s: str) -> str:
    return " ".join(unidecode((s or "")).lower().split())

def to_float(x):
    try: return float(str(x).replace(",", "."))
    except: return None

def normalize_unit(val, unit):
    """Normaliza unidades y devuelve (valor_normalizado, unidad_objetivo)."""
    if val is None: return None, unit
    u = (unit or "").strip().lower()
    v = to_float(val)

    # ppm == mg/kg
    if u in ["ppm", "mg/kg", "mg·kg-1", "mg kg-1"]: return v, "ppm"
    # porcentaje
    if u in ["%", "g/100g"]: return v, "%"
    # temperatura
    if u in ["°f", "f"]: return round((v - 32) * 5/9, 3), "°c"
    if u in ["°c", "c"]: return v, "°c"
    # tiempo
    if u in ["dias", "día", "d", "days"]: return round((v or 0)/30.0, 3), "meses"
    if u in ["mes", "meses", "m", "months", "month"]: return v, "meses"
    # viscosidad
    if u in ["cp", "cps"]: return v, "cP"
    return v, u

# Sinónimos ES/EN para mapear variables
SYN = {
    "vida util": ["vida util","shelf life","expiry","best before","caducidad"],
    "humedad": ["humedad","moisture","water content"],
    "% cacao": ["% cacao","cocoa content","cocoa %","contenido de cacao","% cocoa"],
    "% maltitol": ["% maltitol","maltitol content","polyols","polyols content","polioles"],
    "propiedades fisicoquimicas": ["fisicoquimicas","physicochemical","typical analysis"],
    "propiedades microbiologicas": ["microbiologicas","microbiological"],
    "metales pesados": ["heavy metals","lead","pb","mercury","hg","arsenic","as","cadmium","cd"],
    "aminograma": ["amino acid profile","aminograma","amino acids","aminoacidogram"],
    "alergenos": ["allergens","contains","may contain","alergenos"],
    "gmo": ["gmo","non-gmo","ogm","genetically modified"],
    "almacenamiento": ["almacenamiento","storage","store at","storage conditions"],
    "envase": ["envase","packaging","empaque","container","drum","bag","sack","ibc"],
    "certificaciones": ["haccp","fssc","brc","iso 22000","kosher","halal","certificaciones"],
    "micotoxinas": ["mycotoxins","aflatoxin","ochratoxin","zearalenone","micotoxinas"],
    "plaguicidas": ["pesticides","mrl","rsa","lmrs","plaguicidas"],
    "viscosidad": ["viscosity","cp","centipoise","viscosidad"],
    "punto de fusion": ["melting point","punto de fusion","m.p."],
}

KEYWORDS = set(sum(SYN.values(), [])) | set(SYN.keys())

def std_var(var_raw: str) -> str:
    """Devuelve la variable estandarizada usando SYN o la cadena normalizada si no hay match fuerte."""
    v = nrm(var_raw)
    best, score = None, 0
    for k, alts in SYN.items():
        for a in [k] + alts:
            s = fuzz.partial_ratio(nrm(a), v)
            if s > score:
                score, best = s, k
    return best if score >= 70 else v

def seems_like_var(v):
    """Filtra basura: debe contener letras y alguna palabra clave."""
    v2 = nrm(v)
    if len(v2) < 4: return False
    if not re.search(r"[a-z]", v2): return False
    return any(nrm(k) in v2 for k in KEYWORDS)

# ================== LECTURA DOCX ==================
def read_docx_text_tables(file):
    doc = Document(file)
    text = "\n".join([p.text for p in doc.paragraphs])

    tables = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([c.text.strip() for c in r.cells])
        if rows: tables.append(pd.DataFrame(rows))
    return text, tables, ["docx: ok"]

# ================== LECTURA PDF ==================
def read_pdf_text_tables_plumber(file):
    text_parts, table_dfs, log = [], [], []
    with pdfplumber.open(file) as pdf:
        for i, p in enumerate(pdf.pages, 1):
            txt = p.extract_text() or ""
            text_parts.append(txt)
            found = False
            try:
                for t in p.extract_tables() or []:
                    df = pd.DataFrame(t)
                    if df.shape[1] > 1 and df.dropna(how="all").shape[0] > 1:
                        table_dfs.append(df); found = True
            except: pass
            log.append(f"pdfplumber p{i}: tables={'ok' if found else 'none'}")
    return "\n".join(text_parts), table_dfs, log

def try_camelot(file):
    try:
        import camelot
        logs, tdfs = [], []
        for flavor in ["lattice","stream"]:
            try:
                file.seek(0)
                tables = camelot.read_pdf(file, pages="all", flavor=flavor)
                for t in tables:
                    df = t.df
                    if df.shape[1] > 1 and df.dropna(how="all").shape[0] > 1:
                        tdfs.append(df)
                logs.append(f"camelot {flavor}: {len(tables)} tables")
                if len(tdfs) > 0: break
            except Exception:
                logs.append(f"camelot {flavor}: error")
        return tdfs, logs
    except Exception:
        return [], ["camelot: not available"]

def try_tabula(file):
    try:
        import tabula
        file.seek(0)
        dfs = tabula.read_pdf(file, pages="all", multiple_tables=True, stream=True)
        tdfs = []
        for df in dfs or []:
            if isinstance(df, pd.DataFrame) and df.shape[1] > 1 and df.dropna(how="all").shape[0] > 1:
                tdfs.append(df)
        return tdfs, [f"tabula: {len(tdfs)} tables"]
    except Exception:
        return [], ["tabula: not available or failed"]

def extract_text_tables(uploaded):
    name = uploaded.name.lower()
    if name.endswith(".docx"):
        uploaded.seek(0)
        return read_docx_text_tables(uploaded)
    elif name.endswith(".pdf"):
        uploaded.seek(0)
        t, dfs, log = read_pdf_text_tables_plumber(uploaded)
        if len(dfs) == 0:
            cdfs, clog = try_camelot(uploaded)
            dfs += cdfs; log += clog
        if len(dfs) == 0:
            tdfs, tlog = try_tabula(uploaded)
            dfs += tdfs; log += tlog
        return t, dfs, log
    else:
        uploaded.seek(0)
        return uploaded.read().decode("utf-8", errors="ignore"), [], ["txt: ok"]

# ================== PARSEO AVANZADO ==================
# Regex genérico de valor (fallback)
RE_VAL = re.compile(
    r'(?P<var>[A-Za-zÁÉÍÓÚÜÑa-záéíóúüñ %/°\-\(\)\.\,]+?)'
    r'[:\-\s]{0,3}'
    r'(?P<op><=|≥|<=|>=|<|>|=|≤|≥|entre)?\s*'
    r'(?P<num>\d+(?:[.,]\d+)?)?\s*'
    r'(?P<unit>%|ppm|mg/kg|g/100g|cP|°c|°f|meses|dias|days|month|months)?',
    re.IGNORECASE
)

def parse_from_text(text):
    rows = []
    for m in RE_VAL.finditer(text or ""):
        var = (m.group("var") or "").strip()
        if not seems_like_var(var): 
            continue
        op  = (m.group("op") or "").replace("<=","≤").replace(">=","≥")
        num = m.group("num"); unit = (m.group("unit") or "").lower()
        if not num: 
            continue
        rows.append({"variable_raw": var, "op": op, "val": num, "unit": unit})
    return pd.DataFrame(rows)

def parse_from_tables(dfs):
    rows = []
    for df in dfs:
        if df.empty: continue
        df = df.fillna("").astype(str)
        header = df.iloc[0].tolist()
        if any(h.strip() for h in header):
            df.columns = header; df = df.iloc[1:]

        for _, r in df.iterrows():
            line = " | ".join([str(x) for x in r.tolist()])
            for m in RE_VAL.finditer(line):
                var = (m.group("var") or "").strip()
                if not seems_like_var(var): 
                    continue
                op  = (m.group("op") or "").replace("<=","≤").replace(">=","≥")
                num = m.group("num"); unit = (m.group("unit") or "").lower()
                if not num: 
                    continue
                rows.append({"variable_raw": var, "op": op, "val": num, "unit": unit})
    return pd.DataFrame(rows)

# ======= PARSEO POR ESQUEMA (Param/Min/Target/Max/Unidad) =======
HEADER_MAP = {
    "parametro": ["parametro","parámetro","parameter","item","descripcion","description","test","characteristic"],
    "min": ["min","mín","minimum","lower"],
    "target": ["target","objetivo","typical","nominal"],
    "max": ["max","máx","maximum","upper","limite sup","limit sup"],
    "unidad": ["unidad","unit","units","u.","measure"],
}

def _best_col_index(cols, wanted):
    cols_norm = [nrm(c) for c in cols]
    best_i, best_s = None, 0
    for i, c in enumerate(cols_norm):
        for w in wanted:
            s = fuzz.partial_ratio(w, c)
            if s > best_s:
                best_i, best_s = i, s
    return best_i if best_s >= 70 else None

def extract_table_schema_first(dfs):
    """Devuelve filas normalizadas con columnas mapeadas (parametro, min, target, max, unidad)."""
    out = []
    for df in dfs:
        if df.empty: continue
        df = df.fillna("").astype(str)
        header = df.iloc[0].tolist()
        if any(h.strip() for h in header):
            df.columns = header; df = df.iloc[1:]
        cols = list(df.columns)

        c_param = _best_col_index(cols, HEADER_MAP["parametro"])
        c_min   = _best_col_index(cols, HEADER_MAP["min"])
        c_tgt   = _best_col_index(cols, HEADER_MAP["target"])
        c_max   = _best_col_index(cols, HEADER_MAP["max"])
        c_unit  = _best_col_index(cols, HEADER_MAP["unidad"])

        # Si tenemos Parametro y (Min o Max), procesamos
        if c_param is not None and (c_min is not None or c_max is not None):
            for _, r in df.iterrows():
                row = {
                    "parametro": r.iloc[c_param] if c_param is not None else "",
                    "min": r.iloc[c_min] if c_min is not None else "",
                    "target": r.iloc[c_tgt] if c_tgt is not None else "",
                    "max": r.iloc[c_max] if c_max is not None else "",
                    "unidad": r.iloc[c_unit] if c_unit is not None else "",
                }
                out.append(row)
    return pd.DataFrame(out)

# ======= REGLAS POR VARIABLE (mejor semántica) =======
def default_op_for(var_std, op, vmin, vmax):
    v = nrm(var_std)
    # prioridades por tipo
    if "humedad" in v:
        # si hay target o max, usar ≤; si sólo min, considerar ≥ (raro)
        if vmax is not None: return "≤", None, vmax
        if vmin is not None: return "≥", vmin, None
    if "% cacao" in v or "% maltitol" in v or "polyol" in v:
        if vmin is not None: return "≥", vmin, None
        if vmax is not None: return "≥", vmax, None
    if "viscosidad" in v or "punto de fusion" in v:
        if vmin is not None and vmax is not None: return "entre", vmin, vmax
    if "vida util" in v:
        if vmin is not None: return "≥", vmin, None
        if vmax is not None: return "≥", vmax, None
    # por defecto
    if vmin is not None and vmax is not None and vmin != vmax: return "entre", vmin, vmax
    if vmax is not None: return "≤", None, vmax
    if vmin is not None: return "≥", vmin, None
    return op or "=", vmin, vmax

# ================== BUILD SPEC ==================
def build_spec_from_schema(schema_df):
    """Convierte el DataFrame esquema a la tabla estándar de requisitos."""
    if schema_df.empty:
        return pd.DataFrame(columns=["variable","criterio","unidad_objetivo","op","min","max"])
    rows = []
    for _, r in schema_df.iterrows():
        var = std_var(r.get("parametro", ""))
        if not seems_like_var(var): 
            continue

        # Números y unidad
        vmin = to_float(r.get("min"))
        vmax = to_float(r.get("max"))
        vtg  = to_float(r.get("target"))
        unit = (r.get("unidad") or "").strip().lower()

        # Si no hay min/max pero hay target, usa target como "=", luego se ajusta por regla
        if vmin is None and vmax is None and vtg is not None:
            vmin = vtg; vmax = vtg

        op0, mn0, mx0 = default_op_for(var, None, vmin, vmax)
        # Normalización de unidad basada en el valor relevante
        if op0 == "≤":
            val_u, uni = normalize_unit(mx0, unit)
            criterio = f"≤ {val_u} {uni}" if val_u is not None else f"≤ ? {uni}"
            rows.append({"variable": var, "criterio": criterio, "unidad_objetivo": uni, "op": "≤", "min": None, "max": val_u})
        elif op0 == "≥":
            val_u, uni = normalize_unit(mn0, unit)
            criterio = f"≥ {val_u} {uni}" if val_u is not None else f"≥ ? {uni}"
            rows.append({"variable": var, "criterio": criterio, "unidad_objetivo": uni, "op": "≥", "min": val_u, "max": None})
        elif op0 == "entre":
            vmin_u, uni = normalize_unit(mn0, unit)
            vmax_u, uni2 = normalize_unit(mx0, unit)
            uni = uni or uni2
            criterio = f"entre {vmin_u} y {vmax_u} {uni}"
            rows.append({"variable": var, "criterio": criterio, "unidad_objetivo": uni, "op": "entre", "min": vmin_u, "max": vmax_u})
        else:
            # '=' exacto
            val_u, uni = normalize_unit(mn0, unit)  # mn0==mx0
            criterio = f"= {val_u} {uni}"
            rows.append({"variable": var, "criterio": criterio, "unidad_objetivo": uni, "op": "=", "min": val_u, "max": val_u})

    if not rows:
        return pd.DataFrame(columns=["variable","criterio","unidad_objetivo","op","min","max"])
    out = pd.DataFrame(rows).drop_duplicates(subset=["variable"], keep="first")
    return out

def build_spec_fallback(text, tables):
    """Si no se detecta esquema, usa el parser genérico."""
    df = pd.concat([parse_from_text(text), parse_from_tables(tables)], ignore_index=True)
    if df.empty:
        return pd.DataFrame(columns=["variable","criterio","unidad_objetivo","op","min","max"])

    std, ops, vals, units = [], [], [], []
    for _, r in df.iterrows():
        vstd = std_var(r["variable_raw"])
        if not seems_like_var(vstd):
            continue
        std.append(vstd)
        ops.append(r["op"])
        val, uni = normalize_unit(r["val"], r["unit"])
        vals.append(val); units.append(uni)
    df["variable"] = std; df["op"] = ops; df["val"] = vals; df["unit"] = units

    # Consolidar por variable/unidad
    sp = (df.groupby(["variable","unit"])
            .agg(min=("val","min"), max=("val","max"))
            .reset_index())
    rows = []
    for _, r in sp.iterrows():
        op0, mn0, mx0 = default_op_for(r["variable"], r.get("op",""), r["min"], r["max"])
        unit = r["unit"]
        if op0 == "entre":
            criterio = f"entre {mn0} y {mx0} {unit}"
        elif op0 == "≤":
            criterio = f"≤ {mx0} {unit}"
        elif op0 == "≥":
            criterio = f"≥ {mn0} {unit}"
        else:
            criterio = f"= {mn0} {unit}"
        rows.append({"variable": r["variable"], "criterio": criterio, "unidad_objetivo": unit,
                     "op": op0, "min": mn0 if op0!="≤" else None, "max": mx0 if op0!="≥" else None})
    return pd.DataFrame(rows).drop_duplicates(subset=["variable"], keep="first")

# ================== PROVEEDOR ==================
VALID_UNITS = {"%","ppm","cp","°c","meses","mg/kg"}

def build_provider(docs):
    """Une texto y tablas de varios documentos ([(name, text, tables)]) y devuelve variable/valor/unidad/fuente."""
    rows = []
    for name, text, tables in docs:
        # 1) esquema si existe
        schema = extract_table_schema_first(tables)
        if not schema.empty:
            for _, r in schema.iterrows():
                var = std_var(r.get("parametro",""))
                if not seems_like_var(var): 
                    continue
                vmin = to_float(r.get("min")); vmax = to_float(r.get("max")); vtg = to_float(r.get("target"))
                unit = (r.get("unidad") or "").strip().lower()
                # priorizar max/min/target como valor "reportado"
                cand = vmax if vmax is not None else (vmin if vmin is not None else vtg)
                val, uni = normalize_unit(cand, unit)
                if val is None and not uni: 
                    continue
                rows.append({"variable": var, "valor": val, "unidad": uni, "fuente": name})
        # 2) fallback genérico
        gen = pd.concat([parse_from_text(text), parse_from_tables(tables)], ignore_index=True)
        for _, r in gen.iterrows():
            var = std_var(r["variable_raw"])
            if not seems_like_var(var): 
                continue
            val, uni = normalize_unit(r["val"], r["unit"])
            if val is None and not uni: 
                continue
            rows.append({"variable": var, "valor": val, "unidad": uni, "fuente": name})

    if not rows:
        return pd.DataFrame(columns=["variable","valor","unidad","fuente"])

    prov = pd.DataFrame(rows)
    # preferir filas con unidad válida
    prov["unit_score"] = prov["unidad"].apply(lambda u: 0 if (u or "") in VALID_UNITS else 1)
    prov = prov.sort_values(["variable","unit_score","fuente"]).drop_duplicates(subset=["variable"], keep="first")
    return prov.drop(columns=["unit_score"])

# ================== COMPARACIÓN ==================
def compare(spec, prov):
    out = []
    pmap = {r["variable"]: r for _, r in prov.iterrows()}
    for _, s in spec.iterrows():
        var = s["variable"]; u = s["unidad_objetivo"]; op = s["op"]
        vmin, vmax = s["min"], s["max"]
        prow = pmap.get(var)

        if prow is None:
            out.append({"Variable": var, "Criterio": s["criterio"], "Proveedor": "—", "Unidad": u or "", "Estado": "No informado", "Fuente": ""})
            continue

        pv = prow["valor"]; pu = prow["unidad"]; fuente = prow["fuente"]
        estado = "Revisar"

        if op=="entre" and pv is not None and vmin is not None and vmax is not None:
            if vmin <= pv <= vmax: estado = "Cumple"
            elif pv < vmin: estado = "Supera requisito"
            else: estado = "No cumple"
        elif op=="≤" and pv is not None and vmax is not None:
            estado = "Cumple" if pv <= vmax else "No cumple"
        elif op=="≥" and pv is not None and vmin is not None:
            estado = "Cumple" if pv >= vmin else "No cumple"
        elif op=="=" and pv is not None and vmin is not None:
            estado = "Cumple" if abs(pv - vmin) < 1e-6 else "No cumple"
        elif pv is None:
            estado = "No informado"

        out.append({"Variable": var, "Criterio": s["criterio"], "Proveedor": pv if pv is not None else "—",
                    "Unidad": u or pu or "", "Estado": estado, "Fuente": fuente})
    return pd.DataFrame(out)

# ================== UI: UPLOADS ==================
col1, col2 = st.columns([1,2], gap="large")
with col1:
    st.header("1) Especificación (PDF/DOCX)")
    spec_file = st.file_uploader("Selecciona archivo", type=["pdf","docx"])
with col2:
    st.header("2) Documentos del Proveedor (múltiples)")
    prov_files = st.file_uploader("Selecciona 1..N archivos", type=["pdf","docx"], accept_multiple_files=True)

run = st.button("Comparar")

# ================== PIPELINE ==================
if run:
    if not spec_file:
        st.error("Falta la Especificación.")
        st.stop()
    if not prov_files:
        st.error("Faltan documentos del Proveedor.")
        st.stop()

    # ---- SPEC ----
    st.subheader("Extrayendo Especificación…")
    s_text, s_tables, s_log = extract_text_tables(spec_file)
    with st.expander("Ver log de extracción de la especificación"):
        st.code("\n".join(s_log)[:4000])

    schema_df = extract_table_schema_first(s_tables)
    if not schema_df.empty:
        spec_df = build_spec_from_schema(schema_df)
        st.caption("Esquema detectado: se usó tabla Parametro/Min/Target/Max/Unidad ✅")
    else:
        spec_df = build_spec_fallback(s_text, s_tables)
        st.caption("No se detectó esquema claro; se usó parser genérico ⚠️")

    st.dataframe(spec_df, use_container_width=True)

    # ---- PROVEEDOR ----
    st.subheader("Consolidando Proveedor…")
    docs = []
    p_logs = []
    for f in prov_files:
        t, tbls, lg = extract_text_tables(f)
        p_logs += [f"== {f.name} =="] + lg
        docs.append((f.name, t, tbls))

    with st.expander("Ver log de extracción del proveedor"):
        st.code("\n".join(p_logs)[:4000])

    prov_df = build_provider(docs)
    st.dataframe(prov_df, use_container_width=True)

    # ---- COMPARACIÓN ----
    st.subheader("Comparación")
    comp_df = compare(spec_df, prov_df)

    priority = ["vida util","humedad","% cacao","% maltitol","propiedades microbiologicas","metales pesados","alergenos","gmo","envase","almacenamiento","certificaciones","viscosidad","punto de fusion"]
    comp_df["__ord"] = comp_df["Variable"].apply(lambda v: priority.index(v) if v in priority else 999)
    comp_df = comp_df.sort_values(["__ord","Variable"]).drop(columns="__ord")
    st.dataframe(comp_df, use_container_width=True)

    st.write("**Resumen**")
    st.json(comp_df["Estado"].value_counts().to_dict())

    # ---- DESCARGAS ----
    out = BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as xw:
        spec_df.to_excel(xw, index=False, sheet_name="Checklist")
        prov_df.to_excel(xw, index=False, sheet_name="Proveedor")
        comp_df.to_excel(xw, index=False, sheet_name="Comparacion")
    st.download_button("⬇️ Descargar paquete (Excel)", data=out.getvalue(),
                       file_name="homologacion_pro.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

else:
    st.info("Sube archivos y presiona **Comparar**.")
    
